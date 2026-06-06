import re
import uuid
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import HTTPException
from pypdf import PdfReader
from supabase import Client

BUCKET_NAME = "training-files"
ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
MIME_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


def validate_training_file(filename: str) -> str:
    extension = Path(filename).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Supported file types: PDF, TXT, MD, and DOCX.",
        )
    return extension


def extract_training_file_text(filename: str, content: bytes) -> str:
    extension = validate_training_file(filename)

    if extension in {".txt", ".md"}:
        return content.decode("utf-8", errors="replace").strip()
    if extension == ".pdf":
        return _extract_pdf_text(content)
    return _extract_docx_text(content)


def upload_training_file(
    supabase: Client,
    *,
    user_id: str,
    filename: str,
    content: bytes,
) -> tuple[str, str]:
    extension = validate_training_file(filename)
    request_id = str(uuid.uuid4())
    safe_name = _sanitize_filename(filename)
    storage_path = f"{user_id}/{request_id}/{safe_name}"
    mime_type = MIME_TYPES[extension]

    try:
        supabase.storage.from_(BUCKET_NAME).upload(
            storage_path,
            content,
            file_options={"content-type": mime_type, "upsert": "false"},
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail="Failed to store uploaded file.") from exc

    return request_id, storage_path


def _sanitize_filename(filename: str) -> str:
    stem = Path(filename).stem.strip() or "upload"
    extension = Path(filename).suffix.lower()
    safe_stem = re.sub(r"[^\w.-]+", "-", stem).strip("-._") or "upload"
    return f"{safe_stem}{extension}"


def _extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not read PDF file.") from exc
    return "\n".join(pages).strip()


def _extract_docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise HTTPException(status_code=400, detail="Could not read DOCX file.") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()
